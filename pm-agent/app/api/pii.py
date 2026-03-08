"""PII 마스킹 다운로드 API. 파일을 받아 마스킹 문서 + 대조표 ZIP 반환."""

import csv
import io
import logging
import os
import tempfile
import zipfile

from fastapi import APIRouter, HTTPException, UploadFile
from fastapi.responses import StreamingResponse

from app.services.pii.masker import PIIMasker

logger = logging.getLogger(__name__)

router = APIRouter(prefix="/agent/pii")

ALLOWED_EXTENSIONS = {".pdf", ".docx", ".txt", ".md", ".csv", ".xlsx", ".xls", ".hwp"}
MAX_FILE_SIZE = 5 * 1024 * 1024  # 5MB


def _extract_text(file_path: str) -> str:
    """파일에서 텍스트 추출."""
    ext = os.path.splitext(file_path)[1].lower()

    if ext == ".pdf":
        import fitz
        doc = fitz.open(file_path)
        text = "\n".join(page.get_text() for page in doc)
        doc.close()
        return text
    elif ext == ".docx":
        from docx import Document
        doc = Document(file_path)
        return "\n".join(p.text for p in doc.paragraphs)
    elif ext in (".txt", ".md", ".csv"):
        with open(file_path, "r", encoding="utf-8") as f:
            return f.read()
    elif ext in (".xlsx", ".xls"):
        from openpyxl import load_workbook
        wb = load_workbook(file_path, data_only=True)
        parts = []
        for sheet in wb.sheetnames:
            ws = wb[sheet]
            parts.append(f"[시트: {sheet}]")
            for row in ws.iter_rows(values_only=True):
                cells = [str(c) if c is not None else "" for c in row]
                if any(cells):
                    parts.append("\t".join(cells))
        wb.close()
        return "\n".join(parts)
    elif ext == ".hwp":
        import olefile
        ole = olefile.OleFileIO(file_path)
        parts = []
        if ole.exists("PrvText"):
            parts.append(ole.openstream("PrvText").read().decode("utf-16-le", errors="ignore"))
        elif ole.exists("BodyText/Section0"):
            for entry in ole.listdir():
                if entry[0] == "BodyText":
                    raw = ole.openstream(entry).read()
                    decoded = raw.decode("utf-16-le", errors="ignore")
                    parts.append("".join(c for c in decoded if c.isprintable() or c in "\n\t"))
        ole.close()
        return "\n".join(parts)
    else:
        raise ValueError(f"Unsupported file type: {ext}")


def _apply_mask_to_file(src_path: str, dst_path: str, mapping) -> None:
    """원본 파일 포맷을 유지하면서 PII를 마스킹 토큰으로 치환."""
    ext = os.path.splitext(src_path)[1].lower()

    if not mapping.pii_detected:
        # PII 없으면 원본 그대로 복사
        import shutil
        shutil.copy2(src_path, dst_path)
        return

    # 치환 맵: 원본값 → 마스크 토큰 (긴 것부터 치환하여 부분 매칭 방지)
    replacements = sorted(
        [(orig, token) for token, orig in mapping.mappings.items()],
        key=lambda x: len(x[0]),
        reverse=True,
    )

    def _replace_text(text: str) -> str:
        for orig, token in replacements:
            text = text.replace(orig, token)
        return text

    if ext == ".pdf":
        _mask_pdf(src_path, dst_path, replacements)
    elif ext == ".docx":
        _mask_docx(src_path, dst_path, _replace_text)
    elif ext in (".txt", ".md", ".csv"):
        _mask_text_file(src_path, dst_path, _replace_text)
    elif ext in (".xlsx", ".xls"):
        _mask_xlsx(src_path, dst_path, _replace_text)
    elif ext == ".hwp":
        # HWP는 바이너리 포맷이라 원본 유지 불가 → txt로 fallback
        _mask_text_file(src_path, dst_path, _replace_text, force_text=True)
    else:
        import shutil
        shutil.copy2(src_path, dst_path)


def _mask_pdf(src_path: str, dst_path: str, replacements: list[tuple[str, str]]) -> None:
    """PDF 내 텍스트를 마스킹 토큰으로 치환."""
    import fitz

    doc = fitz.open(src_path)
    for page in doc:
        for orig, token in replacements:
            hits = page.search_for(orig)
            for rect in hits:
                page.add_redact_annot(rect, text=token, fontsize=0)
        page.apply_redactions()
    doc.save(dst_path)
    doc.close()


def _mask_docx(src_path: str, dst_path: str, replace_fn) -> None:
    """DOCX 내 텍스트를 마스킹 토큰으로 치환."""
    from docx import Document

    doc = Document(src_path)
    for para in doc.paragraphs:
        for run in para.runs:
            if run.text:
                run.text = replace_fn(run.text)
    # 테이블 내 텍스트도 처리
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        if run.text:
                            run.text = replace_fn(run.text)
    doc.save(dst_path)


def _mask_text_file(src_path: str, dst_path: str, replace_fn, force_text: bool = False) -> None:
    """텍스트 기반 파일 마스킹."""
    if force_text:
        text = _extract_text(src_path)
    else:
        with open(src_path, "r", encoding="utf-8") as f:
            text = f.read()
    masked = replace_fn(text)
    with open(dst_path, "w", encoding="utf-8") as f:
        f.write(masked)


def _mask_xlsx(src_path: str, dst_path: str, replace_fn) -> None:
    """Excel 파일 내 셀 텍스트를 마스킹 토큰으로 치환."""
    from openpyxl import load_workbook

    wb = load_workbook(src_path)
    for ws in wb.worksheets:
        for row in ws.iter_rows():
            for cell in row:
                if cell.value is not None and isinstance(cell.value, str):
                    cell.value = replace_fn(cell.value)
    wb.save(dst_path)
    wb.close()


@router.post("/mask")
async def mask_document(file: UploadFile):
    """파일을 PII 마스킹하여 마스킹 문서 + 대조표 ZIP으로 반환."""
    filename = file.filename or "unknown"
    ext = os.path.splitext(filename)[1].lower()

    if ext not in ALLOWED_EXTENSIONS:
        raise HTTPException(status_code=400, detail=f"지원하지 않는 파일 형식입니다: {ext}")

    content = await file.read()
    if len(content) > MAX_FILE_SIZE:
        raise HTTPException(status_code=400, detail="파일 크기가 5MB를 초과합니다")

    # 임시 파일로 저장 후 텍스트 추출
    tmp_path = ""
    try:
        with tempfile.NamedTemporaryFile(suffix=ext, delete=False) as tmp:
            tmp.write(content)
            tmp_path = tmp.name

        logger.info("Extracting text: file=%s, ext=%s, size=%d, tmp=%s", filename, ext, len(content), tmp_path)
        text = _extract_text(tmp_path)
        logger.info("Extracted %d chars from %s", len(text), filename)
    except Exception as e:
        logger.error("Text extraction failed: file=%s, error=%s", filename, e, exc_info=True)
        if tmp_path and os.path.exists(tmp_path):
            os.unlink(tmp_path)
        raise HTTPException(status_code=400, detail=f"텍스트 추출 실패: {e}")

    if not text.strip():
        if tmp_path and os.path.exists(tmp_path):
            os.unlink(tmp_path)
        raise HTTPException(status_code=400, detail="파일에서 텍스트를 추출할 수 없습니다")

    # PII 마스킹
    masker = PIIMasker()
    _, mapping = masker.mask(text)

    # 원본 포맷 유지하면서 마스킹 적용
    base_name = os.path.splitext(filename)[0]
    # HWP는 바이너리 포맷이라 txt로 fallback
    out_ext = ".txt" if ext == ".hwp" else ext
    masked_filename = f"{base_name}_masked{out_ext}"

    tmp_out_path = ""
    try:
        with tempfile.NamedTemporaryFile(suffix=out_ext, delete=False) as tmp_out:
            tmp_out_path = tmp_out.name

        _apply_mask_to_file(tmp_path, tmp_out_path, mapping)

        with open(tmp_out_path, "rb") as f:
            masked_bytes = f.read()
    except Exception as e:
        logger.error("Masking failed: file=%s, error=%s", filename, e, exc_info=True)
        raise HTTPException(status_code=500, detail=f"마스킹 처리 실패: {e}")
    finally:
        if tmp_path and os.path.exists(tmp_path):
            os.unlink(tmp_path)
        if tmp_out_path and os.path.exists(tmp_out_path):
            os.unlink(tmp_out_path)

    # ZIP 생성 (마스킹 문서 + 대조표)
    zip_buffer = io.BytesIO()

    with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zf:
        # 마스킹 문서 (원본 포맷 유지)
        zf.writestr(masked_filename, masked_bytes)

        # 대조표 CSV
        csv_buffer = io.StringIO()
        writer = csv.writer(csv_buffer)
        writer.writerow(["토큰", "원본값", "카테고리"])
        if mapping.pii_detected:
            for token, original in mapping.mappings.items():
                category = mapping.categories.get(token, "UNKNOWN")
                writer.writerow([token, original, category])
        else:
            writer.writerow(["-", "(PII 미감지)", "-"])
        zf.writestr(f"{base_name}_pii_mapping.csv", csv_buffer.getvalue())

    zip_buffer.seek(0)

    from urllib.parse import quote

    encoded_name = quote(f"{base_name}_masked.zip")
    return StreamingResponse(
        zip_buffer,
        media_type="application/zip",
        headers={"Content-Disposition": f"attachment; filename*=UTF-8''{encoded_name}"},
    )
